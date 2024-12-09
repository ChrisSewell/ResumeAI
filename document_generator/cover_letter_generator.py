from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from typing import Any
import logging
from pathlib import Path
from datetime import datetime
from agents.models import CoverLetter

class CoverLetterDocumentGenerator:
    """Generates formatted cover letter documents."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def _create_styles(self, doc: Document) -> None:
        """Create consistent styles for the document."""
        # Normal text style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        style.paragraph_format.space_after = Pt(12)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        # Heading style for greeting
        greeting_style = doc.styles.add_style('Greeting', WD_STYLE_TYPE.PARAGRAPH)
        font = greeting_style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        font.bold = True
        greeting_style.paragraph_format.space_after = Pt(12)
        
        # Date style
        date_style = doc.styles.add_style('Date', WD_STYLE_TYPE.PARAGRAPH)
        font = date_style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        date_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_style.paragraph_format.space_after = Pt(12)
        
        # Signature style
        signature_style = doc.styles.add_style('Signature', WD_STYLE_TYPE.PARAGRAPH)
        font = signature_style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        signature_style.paragraph_format.space_before = Pt(24)
        signature_style.paragraph_format.space_after = Pt(0)
    
    def generate(self, cover_letter: CoverLetter, company_name: str, output_path: Path) -> Path:
        """Generate a formatted Word document from cover letter content."""
        doc = Document()
        
        # Create styles
        self._create_styles(doc)
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
        
        # Add date
        date = doc.add_paragraph(style='Date')
        date.add_run(datetime.now().strftime("%B %d, %Y"))
        
        # Add greeting
        greeting = doc.add_paragraph(style='Greeting')
        greeting.add_run(cover_letter.greeting)
        
        # Opening paragraph
        opening = doc.add_paragraph(style='Normal')
        opening.add_run(cover_letter.opening_paragraph)
        
        # Body paragraphs
        for paragraph in cover_letter.body_paragraphs:
            body = doc.add_paragraph(style='Normal')
            body.add_run(paragraph)
        
        # Closing
        closing = doc.add_paragraph(style='Normal')
        closing.add_run(cover_letter.closing_paragraph)
        
        # Signature
        signature_lines = cover_letter.signature.split('\n')
        for i, line in enumerate(signature_lines):
            signature = doc.add_paragraph(style='Signature')
            signature.add_run(line)
            if i == 0:  # Add extra space after the "Sincerely,"
                signature.paragraph_format.space_after = Pt(12)
        
        # Save document with sanitized company name
        safe_company_name = "".join(c for c in company_name if c.isalnum() or c in "_ -").strip()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        doc_path = output_path / f"cover_letter_{safe_company_name}_{timestamp}.docx"
        
        try:
            doc.save(str(doc_path))
            self.logger.info(f"Cover letter document generated: {doc_path}")
        except Exception as e:
            self.logger.error(f"Error saving cover letter: {str(e)}")
            raise
        
        return doc_path 