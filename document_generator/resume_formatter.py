from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Any, Dict, List, Set
import logging
from pathlib import Path
from datetime import datetime
from agents.models import ValidatedResume, JobRequirement

class ResumeDocumentGenerator:
    """Generates formatted Word documents from resume data."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def _format_skills_section(self, skills: Dict[str, List[str]], job_requirements: JobRequirement) -> Dict[str, List[str]]:
        """Format and prioritize skills based on relevance to job requirements."""
        MAX_TECHNICAL_SKILLS = 8
        MAX_SOFT_SKILLS = 5
        MAX_OTHER_SKILLS = 5
        
        def extract_key_terms(job_requirements: JobRequirement) -> Set[str]:
            """Extract key technical terms from job requirements."""
            terms = set()
            
            # Get terms from technical requirements
            for category, skills in job_requirements.technical_requirements.items():
                terms.update(skill.lower() for skill in skills)
            
            # Get terms from responsibilities
            for category, resps in job_requirements.key_responsibilities.items():
                for resp in resps:
                    # Split into words and clean
                    words = set(word.lower().strip('.,()[]') for word in resp.split())
                    # Filter out common words and keep technical/relevant terms
                    technical_terms = {
                        word for word in words 
                        if (
                            len(word) > 2  # Skip very short words
                            and word not in {
                                'the', 'and', 'for', 'with', 'will', 'able', 'must',
                                'can', 'may', 'should', 'would', 'could', 'have', 'has',
                                'had', 'been', 'was', 'were', 'are', 'our', 'your', 'their'
                            }  # Skip common words
                            and not word.isdigit()  # Skip numbers
                        )
                    }
                    terms.update(technical_terms)
            
            # Get terms from qualifications
            for category, quals in job_requirements.required_qualifications.items():
                for qual in quals:
                    words = set(word.lower().strip('.,()[]') for word in qual.split())
                    relevant_terms = {
                        word for word in words 
                        if len(word) > 2 
                        and word not in {'the', 'and', 'for', 'with', 'years', 'year', 'experience'}
                    }
                    terms.update(relevant_terms)
            
            # Add any specific technical terms from soft skills
            for category, soft_skills in job_requirements.soft_skills.items():
                for skill in soft_skills:
                    if any(tech_word in skill.lower() for tech_word in 
                          ['data', 'technical', 'system', 'analysis', 'development']):
                        terms.add(skill.lower())
            
            return terms
        
        def prioritize_skills(skill_list: List[str], required_skills: Set[str], key_terms: Set[str]) -> List[str]:
            """Sort skills by relevance to job requirements."""
            skill_scores = []
            for skill in skill_list:
                score = 0
                skill_lower = skill.lower()
                
                # Direct match with requirement
                if skill_lower in required_skills:
                    score += 100
                # Partial match with requirement
                elif any(req in skill_lower or skill_lower in req for req in required_skills):
                    score += 50
                # Contains key technical terms from job description
                elif any(term in skill_lower for term in key_terms):
                    score += 25
                
                skill_scores.append((skill, score))
            
            # Sort by score and return skills
            return [s[0] for s in sorted(skill_scores, key=lambda x: x[1], reverse=True)]

        # Get required skills and key terms
        required_skills = set()
        for category in job_requirements.technical_requirements.values():
            required_skills.update(skill.lower() for skill in category)
        for category in job_requirements.soft_skills.values():
            required_skills.update(skill.lower() for skill in category)
        
        key_terms = extract_key_terms(job_requirements)
        
        formatted_skills = {}
        
        # Technical skills - prioritize and limit
        if 'technical' in skills:
            tech_skills = prioritize_skills(skills['technical'], required_skills, key_terms)
            formatted_skills['Technical'] = tech_skills[:MAX_TECHNICAL_SKILLS]
        
        # Soft skills - prioritize and limit
        if 'soft' in skills:
            soft_skills = prioritize_skills(skills['soft'], required_skills, key_terms)
            formatted_skills['Professional'] = soft_skills[:MAX_SOFT_SKILLS]
        
        # Other/management skills - prioritize and limit
        if 'other' in skills:
            other_skills = prioritize_skills(skills['other'], required_skills, key_terms)
            formatted_skills['Management'] = other_skills[:MAX_OTHER_SKILLS]
        
        return formatted_skills

    def generate(self, resume_data: ValidatedResume, job_requirements: JobRequirement, output_path: Path) -> Path:
        """Generate a formatted Word document from resume data."""
        doc = Document()
        
        # Set even tighter margins for the document
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.15)     # Reduced from 0.3
            section.bottom_margin = Inches(0.15)   # Reduced from 0.3
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Add name with minimal spacing
        name = doc.add_paragraph()
        name.space_before = Pt(0)
        name.space_after = Pt(2)    # Reduced from 3pt
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name.add_run(resume_data.name)
        name_run.bold = True
        name_run.font.size = Pt(16)
        
        # Add contact info with reduced spacing
        if resume_data.personal_information:
            contact_info = []
            if resume_data.personal_information.contact.get("email"):
                contact_info.append(resume_data.personal_information.contact["email"])
            if resume_data.personal_information.contact.get("phone"):
                phone = resume_data.personal_information.contact["phone"]
                if resume_data.personal_information.contact.get("phone_prefix"):
                    phone = f"{resume_data.personal_information.contact['phone_prefix']}{phone}"
                contact_info.append(phone)
            
            # Add online presence
            if resume_data.personal_information.online_presence:
                for platform, url in resume_data.personal_information.online_presence.items():
                    contact_info.append(url)
            
            if contact_info:
                contact = doc.add_paragraph()
                contact.space_before = Pt(0)
                contact.space_after = Pt(3)    # Reduced from 6pt
                contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
                contact.add_run(" | ".join(contact_info))
        
        # Add summary with minimal spacing
        if resume_data.summary:
            summary_heading = doc.add_heading('Professional Summary', level=1)
            summary_heading.space_before = Pt(6)  # Minimal space before heading
            summary_heading.space_after = Pt(3)   # Minimal space after heading
            summary = doc.add_paragraph(resume_data.summary)
            summary.space_after = Pt(6)  # Reduced spacing after summary
        
        # Add work experience with consistent minimal spacing
        if resume_data.work_experience:
            exp_heading = doc.add_heading('Professional Experience', level=1)
            exp_heading.space_before = Pt(6)
            exp_heading.space_after = Pt(3)
            
            for idx, exp in enumerate(resume_data.work_experience):
                # Determine number of bullet points based on recency
                max_bullets = 5 if idx == 0 else (
                    4 if idx == 1 else (
                    3 if idx == 2 else (
                    2 if idx == 3 else 1
                )))
                
                # Company and Position with minimal spacing
                p = doc.add_paragraph()
                p.space_before = Pt(1)
                p.space_after = Pt(1)
                position_company = f"{exp.position} - {exp.company}"
                p.add_run(position_company).bold = True
                
                # Period and Location
                if exp.employment_period or exp.location:
                    period_loc = []
                    if exp.employment_period:
                        period_loc.append(exp.employment_period)
                    if exp.location:
                        period_loc.append(exp.location)
                    p.add_run(f"\n{' | '.join(period_loc)}")
                
                # Add responsibilities with minimal spacing
                if exp.responsibilities:
                    for i, resp in enumerate(exp.responsibilities[:max_bullets]):
                        bullet = doc.add_paragraph(style='List Bullet')
                        bullet.space_before = Pt(0)
                        bullet.space_after = Pt(0)
                        bullet.add_run(resp)
                        
                        # Remove the paragraph mark spacing
                        bullet._element.get_or_add_pPr().get_or_add_spacing().after = 0
                        
                        # If this is the last bullet point of the current job
                        if i == len(exp.responsibilities[:max_bullets]) - 1:
                            bullet.space_after = Pt(0)  # Ensure no space after last bullet
                            
                            # Only add minimal spacing if not the last job
                            if idx < len(resume_data.work_experience) - 1:
                                bullet.space_after = Pt(1)
        
        # Add skills section with prioritized and limited skills
        if resume_data.skills:
            skills_heading = doc.add_heading('Skills', level=1)
            skills_heading.space_before = Pt(6)
            skills_heading.space_after = Pt(3)
            
            formatted_skills = self._format_skills_section(resume_data.skills, job_requirements)
            
            for category, skills in formatted_skills.items():
                if skills:
                    p = doc.add_paragraph()
                    p.space_before = Pt(0)
                    p.space_after = Pt(2)
                    p.add_run(f"{category}: ").bold = True
                    p.add_run(", ".join(skills))
        
        # Save document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        doc_path = output_path / f"resume_{timestamp}.docx"
        doc.save(str(doc_path))
        self.logger.info(f"Resume document generated: {doc_path}")
        
        return doc_path